"""
Resume parser module — adapted from resume_parser.txt reference.
Uses Google Gemini (not Groq) for consistency with the rest of the platform.
Supports PDF and DOCX resume uploads.
"""

import json
import os
import tempfile
from pathlib import Path

try:
    import google.generativeai as genai
except ImportError:  # local fallback still works without the optional SDK
    genai = None
from dotenv import load_dotenv

from models import ParsedResume, ProgramMatch

load_dotenv()
if genai is not None:
    genai.configure(api_key=os.getenv("GEMINI_API_KEY"))
    model = genai.GenerativeModel("gemini-2.0-flash")
else:
    model = None


# ─── File Reading ─────────────────────────────────────────────

def clean_extracted_text(text: str) -> str:
    """Clean up text, fixing character-spaced PDFs (e.g., 'A s h i s h' -> 'Ashish')."""
    import re
    if not text:
        return ""
    # Detect if line has excessive single-char spaces
    sample = text[:500]
    single_char_spaces = len(re.findall(r'\b\w\s\w\b', sample))
    if single_char_spaces > 10:
        # De-space single letter words
        text = re.sub(r'(?<=\b\w)\s(?=\w\b)', '', text)
        # Fix emails that got spaces around @ or dots
        text = re.sub(r'\s*@\s*', '@', text)
        text = re.sub(r'\s*\.\s*', '.', text)
    return text


def read_pdf(file_path: str) -> str:
    """Extract text from a PDF file with automatic kerning repair."""
    from pypdf import PdfReader
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return clean_extracted_text(text)


def read_docx(file_path: str) -> str:
    """Extract text from a DOCX file."""
    from docx import Document
    document = Document(file_path)
    text = ""
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text += paragraph.text + "\n"
    # Also extract text from tables
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += cell.text + "\n"
    return text


def read_resume_file(file_path: str) -> str:
    """Read text from a PDF or DOCX file."""
    path = Path(file_path)
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return read_pdf(file_path)
    elif suffix == ".docx":
        return read_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}. Only PDF and DOCX are supported.")


# ─── Gemini Resume Parsing ────────────────────────────────────

RESUME_PARSE_PROMPT = """You are an expert resume parser for Indian skill development and vocational education context.

Extract structured information from the following resume text. Return ONLY valid JSON with NO markdown formatting, NO code blocks, NO extra text.

The JSON must have exactly these fields:
{{
  "name": "string or null",
  "email": "string or null",
  "phone": "string or null",
  "total_experience_years": "number or null",
  "skills": ["list of ALL skills mentioned anywhere in the resume"],
  "education": ["list of education entries"],
  "projects": ["list of project titles/descriptions"],
  "certifications": ["list of certifications"],
  "experiences": [
    {{
      "company": "string or null",
      "role": "string or null",
      "duration": "string or null",
      "description": "string or null",
      "skills_used": ["list of skills used"]
    }}
  ]
}}

IMPORTANT RULES:
1. Do NOT invent information — only extract what's actually in the resume
2. Extract skills from ALL sections (skills, experience, projects, certifications)
3. Include internships inside experiences
4. If a value is not available, return null
5. If a list has no information, return an empty list
6. Be thorough — extract every skill, technology, and tool mentioned
7. Return ONLY the JSON object, no markdown, no explanation

RESUME TEXT:
{resume_text}
"""


def _fallback_parse_resume(resume_text: str) -> ParsedResume:
    """Intelligent regex/NLP fallback extractor for resumes when Gemini key is invalid/unavailable."""
    import re
    lines = [line.strip() for line in resume_text.split("\n") if line.strip()]
    text_lower = resume_text.lower()

    # 1. Email extraction
    email_match = re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', resume_text)
    email = email_match.group(0) if email_match else None

    # 2. Phone extraction
    phone_match = re.search(r'(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}|\b\d{10}\b', resume_text)
    phone = phone_match.group(0) if phone_match else None

    # 3. Name extraction (first non-empty line that isn't an email/phone)
    name = "Candidate"
    for line in lines[:5]:
        if not re.search(r'[@\d]', line) and len(line.split()) <= 4 and len(line) > 2:
            name = line.strip()
            break
    if name == "Candidate" and email:
        name = email.split("@")[0].replace(".", " ").title()

    # 4. Comprehensive Skills Database search
    skill_corpus = [
        # Programming & IT
        "Python", "Java", "C++", "C", "JavaScript", "TypeScript", "HTML", "CSS", "React", "Node.js",
        "SQL", "MySQL", "PostgreSQL", "MongoDB", "Git", "GitHub", "Linux", "Docker", "AWS", "Machine Learning",
        "Deep Learning", "Data Analysis", "Pandas", "NumPy", "TensorFlow", "PyTorch", "NLP", "Cybersecurity",
        "Data Structures", "Algorithms", "Object-Oriented Programming", "OOP", "REST API",
        # Core Engineering & Vocational Trades
        "CNC programming", "G-code", "M-code", "Lathe operation", "Milling", "Turner", "Machinist",
        "Blueprint reading", "GD&T", "Precision measurement", "Micrometers", "Vernier calipers", "CMM",
        "AutoCAD", "SolidWorks", "CATIA", "MasterCAM", "ANSYS",
        "PLC programming", "SCADA", "Siemens TIA Portal", "Allen Bradley", "Industrial Automation",
        "Industrial IoT", "IoT sensors", "MQTT", "OPC-UA", "Embedded C", "Arduino", "Raspberry Pi",
        "MIG welding", "TIG welding", "Arc welding", "Gas welding", "Sheet metal fabrication",
        "Lithium-ion battery", "BMS diagnostics", "EV powertrain", "High-voltage safety", "Electric Vehicles",
        "Electrical wiring", "Motor control circuits", "Transformer maintenance", "Power electronics",
        "PCB design", "PCB troubleshooting", "Soldering", "Oscilloscope", "Multimeter",
        "HVAC", "Refrigeration", "Split AC installation", "VRV/VRF systems", "Compressor overhaul",
        "Plumbing", "Pipe fitting", "Hydraulics", "Pneumatics", "Pump maintenance",
        "Solar PV installation", "Inverter setup", "Net metering", "Renewable energy",
        "5S", "Lean manufacturing", "Kaizen", "Quality inspection", "Total Quality Management", "SPC"
    ]

    extracted_skills = []
    for sk in skill_corpus:
        pattern = r'\b' + re.escape(sk.lower()) + r'\b'
        if re.search(pattern, text_lower):
            extracted_skills.append(sk)

    # If few skills matched, also look for items under a 'Skills' section
    if len(extracted_skills) < 4:
        in_skills_section = False
        for line in lines:
            if any(k in line.lower() for k in ["skills", "technical skills", "competencies", "expertise"]):
                in_skills_section = True
                continue
            if in_skills_section:
                if any(k in line.lower() for k in ["education", "experience", "projects", "certifications"]):
                    break
                items = re.split(r'[,|•\-\n]', line)
                for item in items:
                    clean = item.strip()
                    if 2 < len(clean) < 30 and clean not in extracted_skills:
                        extracted_skills.append(clean)

    # 5. Education extraction
    education = []
    for line in lines:
        if any(deg in line.lower() for deg in ["b.tech", "b.e", "bachelor", "diploma", "iti", "m.tech", "master", "hsc", "ssc", "10th", "12th", "cbse", "university", "institute", "college"]):
            if len(line) < 100:
                education.append(line.strip())

    # 6. Certifications
    certifications = []
    for line in lines:
        if any(c in line.lower() for c in ["certified", "certification", "coursera", "nptel", "udemy", "internship", "workshop"]):
            if len(line) < 100 and line not in certifications:
                certifications.append(line.strip())

    return ParsedResume(
        name=name,
        email=email,
        phone=phone,
        total_experience_years=None,
        skills=list(dict.fromkeys(extracted_skills)),
        education=education[:3],
        projects=[],
        certifications=certifications[:3],
        experiences=[]
    )


async def parse_resume_text(resume_text: str) -> ParsedResume:
    """
    Send resume text to Gemini for structured extraction.
    Falls back to intelligent local regex/NLP parser if API key is invalid/unavailable.
    """
    api_key = os.getenv("GEMINI_API_KEY", "")
    if genai is None or not api_key or "PASTE" in api_key or "YOUR_KEY" in api_key or len(api_key) < 15:
        # Use intelligent local fallback immediately
        return _fallback_parse_resume(resume_text)

    prompt = RESUME_PARSE_PROMPT.format(resume_text=resume_text)

    try:
        response = model.generate_content(
            prompt,
            generation_config=genai.GenerationConfig(
                temperature=0,
                max_output_tokens=2048,
            )
        )

        raw_text = response.text.strip()
        if raw_text.startswith("```"):
            lines = raw_text.split("\n")
            raw_text = "\n".join(lines[1:-1])

        parsed = json.loads(raw_text)
        return ParsedResume(**parsed)

    except Exception as e:
        print(f"Gemini resume parsing error ({e}), using fallback parser")
        return _fallback_parse_resume(resume_text)


# ─── Resume → Program Matching ────────────────────────────────

def normalize(text: str) -> str:
    """Normalize for matching."""
    return text.lower().strip()


def skill_matches(skill_a: str, skill_b: str) -> bool:
    """Check if two skills match (fuzzy)."""
    a = normalize(skill_a)
    b = normalize(skill_b)

    # Exact match
    if a == b:
        return True

    # Substring match
    if a in b or b in a:
        return True

    # Word overlap (at least one meaningful word)
    words_a = set(a.split())
    words_b = set(b.split())
    meaningful_overlap = {w for w in (words_a & words_b) if len(w) > 2}
    if meaningful_overlap:
        return True

    return False


def match_resume_to_programs(
    resume: ParsedResume,
    programs: list[dict]
) -> list[ProgramMatch]:
    """
    Compare a parsed resume's skills against all available programs.
    Returns sorted list of program matches with scores.
    """
    results = []

    for program in programs:
        program_skills = program.get("skills_covered", [])
        if not program_skills:
            continue

        matching = []
        missing = []

        for p_skill in program_skills:
            found = False
            for r_skill in resume.skills:
                if skill_matches(r_skill, p_skill):
                    matching.append(p_skill)
                    found = True
                    break
            if not found:
                missing.append(p_skill)

        # Calculate match score
        total = len(program_skills)
        score = round((len(matching) / total) * 100, 1) if total > 0 else 0

        # Generate verdict
        if score >= 70:
            verdict = "Strong match — your skills align well with this program"
        elif score >= 50:
            verdict = "Good fit — some upskilling needed in specific areas"
        elif score >= 30:
            verdict = "Partial match — significant skill development opportunity"
        else:
            verdict = "Growth opportunity — this program would teach you many new skills"

        results.append(ProgramMatch(
            program_name=program.get("program_name", ""),
            program_type=program.get("program_type", ""),
            match_score=score,
            matching_skills=matching,
            missing_skills=missing,
            verdict=verdict,
            current_demand=program.get("current_demand", "medium"),
            future_outlook=program.get("future_outlook", "stable"),
            avg_salary_range=program.get("avg_salary_range", ""),
        ))

    # Sort by match score descending
    results.sort(key=lambda r: r.match_score, reverse=True)
    return results


async def parse_and_match(file_path: str, programs: list[dict]) -> dict:
    """
    Full pipeline: read file → parse resume → match against programs.
    """
    resume_text = read_resume_file(file_path)
    parsed = await parse_resume_text(resume_text)
    matches = match_resume_to_programs(parsed, programs)

    return {
        "resume": parsed.model_dump(),
        "matches": [m.model_dump() for m in matches],
    }
